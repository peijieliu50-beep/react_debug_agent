# -*- coding: utf-8 -*-
"""
tools/file_parser.py
====================
通用文件解析器：采用「**本地预处理 + 纯文本输入**」方案，把多种格式的文件
统一解析为纯文本字符串，再喂给纯文本大模型 —— 不依赖大模型的多模态能力。

设计依据
--------
- DeepSeek 纯文本模型成本更低、Function Calling 更稳定；
- 本地解析兼容性更强、可控，支持的格式也更多。

支持格式
--------
- 代码/文本类：.py .json .txt .md .csv .log .yaml .yml  → 直接读取
- 文档类      ：.pdf（PyPDF2） .docx（python-docx）     → 提取全文
- 图片类      ：.png .jpg .jpeg .bmp（pytesseract OCR） → 识别报错/曲线截图文字

所有第三方解析库都做了**异常降级**：库未安装或解析失败时返回友好提示，不崩溃。
"""

from pathlib import Path
from typing import Optional

# 各类扩展名
_TEXT_EXT = {".py", ".txt", ".md", ".json", ".csv", ".log", ".yaml", ".yml", ".ini", ".cfg"}
_PDF_EXT = {".pdf"}
_DOCX_EXT = {".docx"}
_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}

SUPPORTED_EXT = _TEXT_EXT | _PDF_EXT | _DOCX_EXT | _IMAGE_EXT


def _clean_text(text: str) -> str:
    """清洗文本：去除多余空白行与首尾空格，保持段落结构。"""
    if not text:
        return ""
    lines = [ln.rstrip() for ln in text.splitlines()]
    # 合并连续空行为最多一个
    cleaned = []
    blank = False
    for ln in lines:
        if ln.strip() == "":
            if not blank:
                cleaned.append("")
            blank = True
        else:
            cleaned.append(ln)
            blank = False
    return "\n".join(cleaned).strip()


# ------------------------------------------------------------
# 各格式解析
# ------------------------------------------------------------
def _parse_text(fp: Path) -> str:
    return fp.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(fp: Path) -> str:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "[解析降级] 未安装 PyPDF2，无法解析 PDF。请 pip install PyPDF2"
    try:
        reader = PdfReader(str(fp))
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            txt = page.extract_text() or ""
            if txt.strip():
                pages.append(f"--- 第 {i} 页 ---\n{txt}")
        if not pages:
            return "[提示] PDF 未提取到文本（可能为扫描件，建议转图片走 OCR）"
        return "\n\n".join(pages)
    except Exception as e:  # noqa: BLE001
        return f"[解析失败] PDF 读取异常: {type(e).__name__}: {e}"


def _parse_docx(fp: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "[解析降级] 未安装 python-docx，无法解析 .docx。请 pip install python-docx"
    try:
        doc = docx.Document(str(fp))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # 解析表格
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        if not parts:
            return "[提示] .docx 文档为空"
        return "\n".join(parts)
    except Exception as e:  # noqa: BLE001
        return f"[解析失败] docx 读取异常: {type(e).__name__}: {e}"


def _parse_image_ocr(fp: Path) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "[解析降级] 未安装 pytesseract / Pillow，无法 OCR。请 pip install pytesseract Pillow，并安装 Tesseract 引擎"

    from config.config import CONFIG
    # Windows 下可能需指定 Tesseract 引擎路径
    if CONFIG.tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = CONFIG.tesseract_cmd
    try:
        img = Image.open(str(fp))
        text = pytesseract.image_to_string(img, lang=CONFIG.ocr_lang)
        if not text.strip():
            return "[提示] OCR 未识别到文字（图片可能无文本或清晰度不足）"
        return text
    except pytesseract.TesseractNotFoundError:
        return "[解析降级] 未找到 Tesseract OCR 引擎，请安装并在 .env 配置 TESSERACT_CMD 路径"
    except Exception as e:  # noqa: BLE001
        return f"[解析失败] OCR 异常: {type(e).__name__}: {e}"


# ------------------------------------------------------------
# 统一入口
# ------------------------------------------------------------
def parse_file(path: str) -> str:
    """把任意支持格式的文件解析为纯文本（已清洗）。

    Args:
        path: 文件路径

    Returns:
        纯文本字符串；失败/降级时返回以 [..] 开头的提示信息。
    """
    fp = Path(path)
    if not fp.exists():
        return f"[错误] 文件不存在: {path}"
    if not fp.is_file():
        return f"[错误] 不是文件: {path}"

    ext = fp.suffix.lower()
    if ext in _TEXT_EXT:
        raw = _parse_text(fp)
    elif ext in _PDF_EXT:
        raw = _parse_pdf(fp)
    elif ext in _DOCX_EXT:
        raw = _parse_docx(fp)
    elif ext in _IMAGE_EXT:
        raw = _parse_image_ocr(fp)
    else:
        return f"[不支持] 暂不支持的文件类型: {ext}。支持: {', '.join(sorted(SUPPORTED_EXT))}"

    # 降级/错误提示直接返回，不再清洗
    if raw.startswith("["):
        return raw
    return _clean_text(raw)


if __name__ == "__main__":
    import tempfile, os
    # 自检：文本与 csv
    d = tempfile.mkdtemp()
    p = Path(d) / "t.py"
    p.write_text("import torch\n\n\n\nprint('hi')\n", encoding="utf-8")
    print("解析 .py:")
    print(parse_file(str(p)))
    print("\n支持的格式:", ", ".join(sorted(SUPPORTED_EXT)))
